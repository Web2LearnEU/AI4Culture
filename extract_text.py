import os
import docx
import textract
import pdfminer.high_level
import pandas as pd

# Folder containing nomination forms
input_folder = "nomination_forms"
output_csv = "extracted_text.csv"
file_data = []
MAX_CHAR_PER_CELL = 32000  # Prevents CSV truncation

def extract_text_from_docx(file_path):
    try:
        doc = docx.Document(file_path)
        text = []

        # Extract text from normal paragraphs
        for para in doc.paragraphs:
            text.append(para.text.strip())

        # Extract text from tables (first column as possible headers)
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    text.append(" | ".join(row_text))

        return "\n".join(text)
    except Exception as e:
        return f"Error extracting DOCX: {e}"

def extract_text_from_doc(file_path):
    try:
        return textract.process(file_path).decode("utf-8").strip()
    except Exception as e:
        return f"Error extracting DOC: {e}"

def extract_text_from_pdf(file_path):
    try:
        return pdfminer.high_level.extract_text(file_path).strip()
    except Exception as e:
        return f"Error extracting PDF: {e}"

# Process each file in the folder
for file_name in os.listdir(input_folder):
    file_path = os.path.join(input_folder, file_name)

    if file_name.endswith(".docx"):
        extracted_text = extract_text_from_docx(file_path)
        file_type = "DOCX"
    elif file_name.endswith(".doc"):
        extracted_text = extract_text_from_doc(file_path)
        file_type = "DOC"
    elif file_name.endswith(".pdf"):
        extracted_text = extract_text_from_pdf(file_path)
        file_type = "PDF"
    else:
        print(f"Skipping unsupported file: {file_name}")
        continue

    # Split long text into smaller chunks
    text_parts = [extracted_text[i:i+MAX_CHAR_PER_CELL] for i in range(0, len(extracted_text), MAX_CHAR_PER_CELL)]
    
    for part_number, text_chunk in enumerate(text_parts):
        file_data.append({
            "File Name": file_name,
            "File Type": file_type,
            "Part": part_number + 1,  # Indicate the part number
            "Extracted Text": text_chunk
        })

# Convert to DataFrame and save
df = pd.DataFrame(file_data)
df.to_csv(output_csv, index=False)

print(f"✅ Text extraction complete! Results saved to '{output_csv}'")

# ------------------------------------
# 📌 Step 2: Load and Analyze Extracted Data
# ------------------------------------

# Load the extracted text
df = pd.read_csv(output_csv)

# Basic statistics
num_files = df["File Name"].nunique()
num_parts = df.shape[0]
avg_text_length = df["Extracted Text"].str.len().mean()

print("\n📊 Basic Dataset Info:")
print(f"Total Unique Files: {num_files}")
print(f"Total Text Parts: {num_parts}")
print(f"Average Text Length per Part: {round(avg_text_length, 2)} characters")

# Show a sample of extracted text
print("\n📌 Sample Extracted Text:")
print(df[["File Name", "Part", "Extracted Text"]].head(5))
