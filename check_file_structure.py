import os
import pdfminer.high_level
import docx
import textract
import pandas as pd
from thefuzz import fuzz, process  # Fuzzy matching

# Folder containing the nomination forms
input_folder = "nomination_forms"
output_csv = "nomination_structure_summary.csv"

# Expected sections with possible variations
expected_sections = {
    "State(s) Party(ies)": ["State(s) Party(ies)", "State Party", "Submitting State(s)"],
    "Name of Element": ["Name of Element", "Element Name", "Title of the Element","Name of the element"],
    "Community(ies), group(s) or, if applicable, individual(s) concerned": [
        "Community(ies), group(s) or, if applicable, individual(s) concerned",
        "Community Concerned", "Involved Communities", "Community Groups"
    ],
    "Brief textual description of the nominated element": [
        "Brief textual description of the nominated element",
        "Brief Description", "Short Description", "Textual Summary"
    ],
    "Brief statement of the viability of the element, its need for safeguarding and the proposed safeguarding measures": [
        "Brief statement of the viability of the element, its need for safeguarding and the proposed safeguarding measures",
        "Viability Statement", "Safeguarding Need", "Conservation Urgency"
    ],
    "Identification of the Element": ["Identification of the Element", "Element Identification"],
    "Name of element": ["Name of element", "Element Name"],
    "Other name(s) of the element, if any": [
        "Other name(s) of the element, if any",
        "Alternative Names", "Other Names"
    ],
    "Identification of the community(ies), group(s) or, if applicable, individual(s) concerned and their location": [
        "Identification of the community(ies), group(s) or, if applicable, individual(s) concerned and their location",
        "Community Identification", "Community Location"
    ],
    "Geographic location and range of the element": [
        "Geographic location and range of the element",
        "Geographic Location", "Element Location", "Range and Area"
    ],
    "Domain(s) represented by the element:": [
        "Domain(s) represented by the element:",
        "Cultural Domains", "Heritage Domains"
    ],
    "Description of the element": ["Description of the element", "Element Description"],
    "Need for urgent safeguarding": [
        "Need for urgent safeguarding",
        "Urgent Safeguarding",
        "Conservation Need", "Protection Requirement"
    ],
    "Viability assessment": ["Viability assessment", "Sustainability Evaluation", "Element Viability"],
    "Threat and risk assessment": [
        "Threat and risk assessment",
        "Risk and Threats",
        "Endangerment Assessment"
    ],
    "Safeguarding measures": ["Safeguarding measures", "Preservation Strategies"],
    "Current and recent efforts to safeguard the element": [
        "Current and recent efforts to safeguard the element",
        "Recent Conservation Efforts",
        "Past Safeguarding Actions"
    ],
    "Safeguarding measures proposed": [
        "Safeguarding measures proposed",
        "Planned Safeguarding Efforts",
        "Future Conservation Plans"
    ],
    "Commitments of States and of communities, groups or individuals concerned": [
        "Commitments of States and of communities, groups or individuals concerned",
        "State and Community Pledges"
    ],
    "Community involvement and consent": [
        "Community involvement and consent",
        "Community Participation",
        "Stakeholder Involvement"
    ],
    "Participation of communities, groups and individuals": [
        "Participation of communities, groups and individuals",
        "Local Engagement",
        "Community Efforts"
    ],
    "Free, prior and informed consent": [
        "Free, prior and informed consent",
        "Informed Consent",
        "Community Consent"
    ],
    "Respect for customary practices governing access": [
        "Respect for customary practices governing access",
        "Customary Practices",
        "Traditional Governance of Access"
    ],
    "Inclusion on an inventory": [
        "Inclusion on an inventory",
        "Listed in Inventory",
        "Recorded in Heritage Inventory"
    ],
    "Documentation": ["Documentation", "Supporting Documents"],
    "Required and supplementary documentation": [
        "Required and supplementary documentation",
        "Supplementary Files",
        "Attached Materials"
    ],
    "Cession of rights": ["Cession of rights", "Transfer of Rights"],
    "List of additional resources": [
        "List of additional resources",
        "Additional References",
        "Extra Materials"
    ],
    "Contact Information": ["Contact Information", "Point of Contact"],
    "Submitting State(s) Party(ies)": ["Submitting State(s) Party(ies)", "State Submitting"],
    "Contact person for correspondence": [
        "Contact person for correspondence",
        "Correspondence Contact",
        "Primary Contact"
    ],
    "Competent body involved": [
        "Competent body involved",
        "Responsible Organization",
        "Involved Authority"
    ],
    "Concerned community organization(s) or representative(s)": [
        "Concerned community organization(s) or representative(s)",
        "Community Representatives"
    ],
    "Signature on behalf of the State(s) Party(ies)": [
        "Signature on behalf of the State(s) Party(ies)",
        "Official Signature"
    ]
}

# Function to find best match using fuzzy matching
def find_best_match(text, section_variations):
    for variation in section_variations:
        if fuzz.partial_ratio(text.lower(), variation.lower()) > 85:  # 85% match threshold
            return True
    return False

def extract_text_from_docx(file_path):
    try:
        doc = docx.Document(file_path)
        text = []

        # Extract text from normal paragraphs
        for para in doc.paragraphs:
            text.append(para.text.strip())

        # Extract text from tables (important!)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text.append(cell.text.strip())  # Extract text from table cells

        return "\n".join(text).strip()
    except Exception as e:
        return f"Error extracting DOCX: {e}"


# Analyze all files
file_data = []
for file_name in os.listdir(input_folder):
    file_path = os.path.join(input_folder, file_name)

    # Extract text based on file type
    if file_name.endswith(".pdf"):
        text = pdfminer.high_level.extract_text(file_path)
        file_type = "PDF"
    elif file_name.endswith(".docx"):
        text = extract_text_from_docx(file_path) 
        file_type = "DOCX"
    elif file_name.endswith(".doc"):
        text = textract.process(file_path).decode("utf-8")
        file_type = "DOC"
    else:
        print(f"Skipping unsupported file: {file_name}")
        continue

    # Compute metadata
    num_words = len(text.split()) if text else 0
    num_lines = len(text.split("\n")) if text else 0

    # Check for expected sections using fuzzy matching
    section_presence = {section: "Yes" if find_best_match(text, variations) else "No" for section, variations in expected_sections.items()}

    # Store data
    file_data.append({
        "File Name": file_name,
        "File Type": file_type,
        "Word Count": num_words,
        "Line Count": num_lines,
        **section_presence  # Expands to include all section checks
    })

# Convert to DataFrame and save
df = pd.DataFrame(file_data)
df.to_csv(output_csv, index=False, encoding="utf-8")

print(f"✅ File structure analysis complete! Results saved to {output_csv}")
